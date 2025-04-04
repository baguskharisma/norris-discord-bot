import os
import logging
from typing import Optional

# initialize logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def parse_document(file_path: str, file_extension: str)->Optional[str]:
    """
    Parse a document and extract its text content.
    
    Args:
        file_path (str): Path to the document file
        file_extension (str): Extension of the file (pdf, docx, txt)
        
    Returns:
        Optional[str]: Extracted text content or None if extraction fails
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        
        return None
    
    try:
        if file_extension == 'pdf':
            return _parse_pdf(file_path)
        elif file_extension == 'docx':
            return _parse_docx(file_path)
        elif file_extension == 'txt':
            return _parse_txt(file_path)
        else:
            logger.error(f"Unsupported file extension: {file_extension}")
            
            return None
    
    except Exception as e:
        logger.error(f"Error parsing document: {str(e)}", exc_info=True)
        
        return None
    
def _parse_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_path (str): Path to the PDF file
    
    Returns:
        str: Extracted text content
    """
    
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            # extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
            return text
    
    except ImportError:
        logger.error("PyPDF2 is not installed.")
        raise
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def _parse_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
    
    Returns:
        str: Extracted text content
    """
    
    try:
        import docx
        
        doc = docx.Document(file_path)
        text = ""
        
        # extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
    
    except ImportError:
        logger.error("python-docx is not installed.")
        raise
    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise

def _parse_txt(file_path: str) -> str:
    """
    Extract text from a TXT file.
    
    Args:
        file_path (str): Path to the TXT file
        
    Returns:
        str: Extracted text content
    """
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    except UnicodeDecodeError:
        # try different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        
        except Exception as e:
            logger.error(f"Error reading TXT file with latin-1 encoding: {str(e)}")
            raise
        
        except Exception as e:
            logger.error(f"Error reading TXT file; {str(e)}")
            raise