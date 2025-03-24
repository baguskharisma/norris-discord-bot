import os
import io
import discord
from discord.ext import commands
from groq import Groq
import docx
from docx import Document
import PyPDF2
import csv

# konfigurasi bot discord dan groq
TOKEN = ""
GROQ_API_KEY = ""
MODEL = "llama-3.3-70b-versatile"

# inisialisasi bot dan klien groq
intents = discord.Intents.default()
bot = commands.Bot(command_prefix='!', intents=intents)
groq_client = Groq(api_key=GROQ_API_KEY)

@bot.event
async def on_ready():
    print(f'{bot.user} already connected to Discord!')
    
    # sinkronisasi perintah "/" commands dengan discord
    try:
        synced = await bot.tree.sync()
        print(f"Sync {len(synced)} command(s)")
        
    except Exception as e:
        print(f"Error sync commands: {e}")

@bot.tree.command(name="summarize", description="Summarize document")

async def slash_summarize(interaction: discord.Interaction, file: discord.Attachment = None):
    """perintah "/" untuk meringkas dokumen"""
    if not file:
        await interaction.response.send_message("Please attach the document to be summarized! Use the command `/summarize` again with the file attachment.", ephemeral=True)
        
        return
    
    # respon langsung untuk mencegah timeout
    await interaction.response.defer(thinking=True)
    
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    if file_extension not in ['.txt', '.pdf', '.docx', '.csv']:
        await interaction.followup.send("Format not supported. Supported format: .txt, .pdf, .docx, .csv")
        
        return
    
    # download file
    file_content = await file.read()
    
    # ekstrak teks berdasarkan tipe file
    extracted_text = extract_text(file_content, file_extension)
    
    if not extracted_text:
        await interaction.followup.send("Failed to extract text from the file.")
        
        return
    
    # ringkas teks menggunakan groq
    summary = await summarize_with_groq(extracted_text)
    
    if not summary:
        await interaction.followup.send("Failed to extract document.")
       
        return
    
    # membuat file dengan ringkasan dalam format yang sama
    output_file = create_output_file(summary, file.filename, file_extension)
    
    # kirim file yang sudah diringkas
    await interaction.followup.send("Here is the summary of your document:", file=discord.File(output_file, f"summary_{file.filename}"))

def extract_text(file_content, file_extension):
    """mengekstrak teks dari berbagai format file"""
    try:
        if file_extension == '.txt':
            return file_content.decode('utf-8')
        
        elif file_extension == '.pdf':
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text
        
        elif file_extension == '.docx':
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            return text
        
        elif file_extension == '.csv':
            csv_file = io.BytesIO(file_content)
            text = ""
            csv_reader = csv.reader(io.StringIO(csv_file.read().decode('utf-8')))
            for row in csv_reader:
                text += ", ".join(row) + "\n"
            
            return text
        
        return None
    
    except Exception as e:
        print(f"Error while extracting file: {e}")
        
        return None

async def summarize_with_groq(text):
    """meringkas teks menggunakan groq api"""
    try:
        # batasi input ke ukuran yang wajar untuk model
        if len(text) > 30000:
            text = text[:30000] + "...[text truncated because it's too long]"
        
        prompt = f"""
        Summarize the following text clearly and concisely. Retain the key points and main information.
        
        Text:
        {text}
        
        Summary:
        """
        
        completion = groq_client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": "You are an assistant who summarizes documents briefly, clearly, and accurately."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2048
        )
        
        return completion.choices[0].message.content
    
    except Exception as e:
        print(f"Error while summarizing text: {e}")
        
        return None

def create_output_file(summary, original_filename, file_extension):
    """membuat file output dalam format yang sama dengan file asli"""
    try:
        # membuat nama file output
        output_filename = f"summary_{original_filename}"
        output_file = io.BytesIO()
        
        if file_extension == '.txt':
            output_file.write(summary.encode('utf-8'))
        
        elif file_extension == '.pdf':
            # perbaikan untuk pembuatan PDF
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.units import inch
            
            # gunakan SimpleDocTemplate untuk membuat PDF yang lebih stabil dan mudah dikelola
            doc = SimpleDocTemplate(
                output_file,
                pagesize=letter,
                rightMargin=72, leftMargin=72,
                topMargin=72, bottomMargin=72
            )
            
            # buat daftar elemen yang akan dirender
            elements = []
            styles = getSampleStyleSheet()
            
            # tambahkan judul
            title_text = f"summary {original_filename}"
            elements.append(Paragraph(title_text, styles['Title']))
            elements.append(Paragraph("<br/>", styles['Normal']))  # spasi
            
            # pecah summary menjadi paragraf
            paragraphs = summary.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    elements.append(Paragraph(para.replace('\n', '<br/>'), styles['Normal']))
                    elements.append(Paragraph("<br/>", styles['Normal']))  # spasi antar paragraf
            
            # build PDF
            doc.build(elements)
        
        elif file_extension == '.docx':
            doc = Document()
            doc.add_heading(f'Summary {original_filename}', 0)
            
            paragraphs = summary.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
            
            doc.save(output_file)
        
        elif file_extension == '.csv':
            writer = csv.writer(output_file)
            writer.writerow(["Summary"])
            writer.writerow([summary])
        
        output_file.seek(0)
        
        return output_file
    
    except Exception as e:
        print(f"Error while creating output file: {e}")
        fallback = io.StringIO(summary)
        fallback.name = f"summary_{original_filename}.txt"
        
        return fallback
    
@bot.command(name="answer", description="Answer questions from the sent document")
async def slash_answer(interaction: discord.Interaction, file: discord.Attachment = None):
    """perintah "/" untuk menjawab pertanyaan dari dokumen apapun"""
    if not file:
        await interaction.response.send_message(
            "Please attach a document containing questions! Use `\answer` again with the file attachment.",
            ephemeral=True
        )
        
        return
    
    await interaction.response.defer(thinking=True)
    
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    if file_extension not in ['.txt', '.pdf', '.docx']:
        await interaction.followup.send("Unsupported format. Supported formats: .txt, .pdf, .docx")
        
        return
    
    # download file
    file_content = await file.read()
    
    #esktrak reks dari file
    extracted_text = extract_text(file_content, file_extension)
    
    if not extracted_text:
        await interaction.followup.send("Failed to extract text from the file.")
        
        return
    
    # kirim teks ke groq api untuk dijawab
    answer = await answer_questions(extracted_text)
    
    if not answer:
        await interaction.followup.send("Failed to generate an answer.")
        
        return
    
    # buat file .txt untuk jawaban
    output_file = io.BytesIO()
    output_file.write(answer.encode('utf-8'))
    output_file.seek(0)
    
    # kirim file jawaban
    await interaction.followup.send(
        "Here are the answers from the document:",
        file=discord.File(output_file, f"answer_{file.filename.split('.')[0]}.txt")
    )
    
def extract_text(file_content, file_extension):
    """mengekstrak teks dari berbagai format file"""
    try:
        if file_extension == '.txt':
            return file_content.decode('utf-8')
        
        elif file_extension == '.pdf':
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            
            return text
        
        return None
    
    except Exception as e:
        print(f"Error extracting file: {e}")
        
        return None

async def answer_questions(text):
    """menjawab pertanyaan dari dokumen menggunakan groq api"""
    try:
        if len(text) > 30000:
            text = text[:30000] + "...[text truncated because it's too long]"
            
        prompt = f"""
        You are an AI assistant that answersnvarious questions from documents. Please analyze the document below and provide the best answers in a clear and structured format.
        
        Document:
        {text}
        
        Answer:
        """
        
        completion = groq_client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": "You are a highly intelligent AI that answers any type of question accurately."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=2048
        )
        
        answer = completion.choices[0].message.content
        
        formatted_answer = "Answers\n"
        formatted_answer += "========================\n\n"
        formatted_answer += answer
        formatted_answer += "\n\n========================\n"
        formatted_answer += "Generated by Norris"
        
        return formatted_answer
    
    except Exception as e:
        print(f"Error while answering: {e}")
        
        return None
    
# jalankan bot
bot.run(TOKEN)